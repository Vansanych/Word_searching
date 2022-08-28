import docx

doc = docx.Document(r'C:\Users\Иван\Desktop\Иван\python\pythonProject\python-docx\01-ТК 001-ППР001-08-2020 Геодез.docx')

target_cell = doc.tables[1].rows[2].cells[0]
target_text = target_cell.paragraphs[0].runs[1].text

target_cell.paragraphs[0].runs[1].text = str('Какой-то текст 4')
print(target_cell.paragraphs[0].runs[1].text)


# doc.save(r'C:\Users\Иван\Desktop\Иван\python\pythonProject\python-docx\01-ТК 001-ППР001-08-2020 Геодез1.docx')

tables = doc.tables
rows1 = tables[1].rows
cells0 = rows1[0].cells

print(len(tables))
print(len(rows1))
print(len(cells0))

tables_text = []
for row in range(len(rows1)):  # создает список из текста в ячейках
    for cell in range(len(rows1[row].cells)):
        tables_text.append(rows1[row].cells[cell].text)


tables_word = []  # создание списка из слов в ячейках
for row in range(len(doc.tables[1].rows)):
    for cell in range(len(doc.tables[1].rows[row].cells)):
        for word in range(len(doc.tables[1].rows[row].cells[cell].text.split())):
            if doc.tables[1].rows[row].cells[cell].text.split()[word] not in tables_word:
                tables_word.append(doc.tables[1].rows[row].cells[cell].text.split()[word])


print(tables_text)
print(tables_word)


tables_letter = ''  # создает строку из слов
for word in tables_word:
    tables_letter += word


letter_list = list(tables_letter)  # разбивает строку из слов по буквам


def searching(letters):

    def search(i):  # выполняет поиск буквы и переходит к следующей
        if i < len(list(letters)):
            if letters[i] in letter_list:
                search(i+1)
            else:
                print('нет такого сочетания букв')
        else:
            print('В тексте найдено слово (' + letters + ')')

    search(0)


searching('Пунга')
